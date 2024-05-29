import math
from random import shuffle
from time import time
import tracemalloc
import matplotlib.pyplot as plt

import numpy as np
import torch
from torch.autograd import Variable


def generate_sample(X, f, size, radius):
    sample_X = []
    sample_Y = []

    min_X = min(X)
    max_X = max(X)

    while len(sample_X) < size:
        x = np.random.uniform(min_X, max_X)
        y = np.random.uniform(f(x) - radius, f(x) + radius)

        sample_X.append(x)
        sample_Y.append(y)

    return np.array(sample_X, dtype=np.float32), np.array(sample_Y, dtype=np.float32)

def print_sample(X, real_Y, found_Y, sample_X, sample_Y):
    fig, ax = plt.subplots()

    ax.scatter(sample_X, sample_Y, color='gray', alpha=0.5, s=10, antialiased=True)
    ax.plot(X, real_Y, label='Real', color='lime', antialiased=True, linewidth=3)
    ax.plot(X, found_Y, label='Found', color='red', antialiased=True, linewidth=2)

    plt.legend()
    plt.show()

#################################################################################

class LinearRegression:
    # first parameter, last factor are constant
    # reg = a + b0x0 + b1x1 + b2x2 ..., point = (x0, x1, x2..., y)
    def diff(reg, point):
        return reg[0] - point[-1] + reg[1:] @ point[0:-1]

    def loss(reg, point):
        return LinearRegression.diff(reg, point)**2

    def average_loss(reg, sample, subsample):
        losses = [LinearRegression.loss(reg, sample[i]) for i in subsample]
        return sum(losses) / len(losses)

    def gradient(reg, point):
        d = LinearRegression.diff(reg, point)
        gr = [2*d]
        for p in point[:-1]:
            gr.append(2*p*d)
        return np.array(gr)

    def average_gradient(reg, sample, subsample):
        grads = [LinearRegression.gradient(reg, sample[i]) for i in subsample]
        return sum(grads) / len(grads)

class PolynomialRegression:
    # first parameter, last factor are constant
    # reg = a0 + a1x + a2x**2 + a3x**3 ..., point = (x, y)
    def diff(reg, point):
        x, y = point[0], point[1]
        y_r = sum(reg[i] * x**i for i in range(len(reg)))
        return y_r - y

    def loss(reg, point):
        return PolynomialRegression.diff(reg, point)**2

    def average_loss(reg, sample, subsample):
        losses = [PolynomialRegression.loss(reg, sample[i]) for i in subsample]
        return sum(losses) / len(losses)

    def gradient(reg, point):
        d = PolynomialRegression.diff(reg, point)
        gr = [2*d * point[0]**i for i in range(len(reg))]
        return np.array(gr)

    def average_gradient(reg, sample, subsample):
        grads = [PolynomialRegression.gradient(reg, sample[i]) for i in subsample]
        return sum(grads) / len(grads)

class LearningRate:
    def constant(value):
        def f(*args):
            return value
        return f
    
    def exponential(start, decay):
        def f(step, *args):
            return start * math.exp(-decay * step)
        return f
    
    def step_based(start, decay, drop_rate):
        def f(step, *args):
            return start * decay**math.floor((step + 1) / drop_rate)
        return f

class L1:
    def __init__(self, importance):
        self.importance = importance

    def func(self, reg):
        return self.importance * sum(np.abs(reg))
    
    def gradient(self, reg):
        return np.array([self.importance for _ in range(len(reg))])

class L2:
    def __init__(self, importance):
        self.importance = importance

    def func(self, reg):
        return self.importance * sum(np.square(reg))
    
    def gradient(self, reg):
        return np.array([self.importance * reg[i] for i in range(len(reg))])

class Elastic:
    def __init__(self, a):
        self.a = a
        self.l1 = L1(a)
        self.l2 = L2(1 - a)
    
    def func(self, reg):
        return self.l1.func(reg) + self.l2.func(reg)
    
    def gradient(self, reg):
        return self.l1.gradient(reg) + self.l2.gradient(reg)

def stochastic_gradient_descent(reg_start, sample, reg_type, learning_rate_func, regularization, batch_size=1, decay=0.5, N=1000, eps=1e-6):
    tracemalloc.start()
    start_time = time()
    iterations = 0
    loss_trace = []

    reg = reg_start

    avg_loss = reg_type.average_loss(reg, sample, range(len(sample)))
    avg_loss_pr = None

    loss_trace.append(avg_loss)

    for epoech in range(N):
        learning_rate = learning_rate_func(epoech)

        batch_order = list(range(len(sample)))
        shuffle(batch_order)

        for b in range(0, max(1, math.floor(len(sample) / batch_size))):
            iterations += 1

            b_begin = batch_size * b
            b_end = min(batch_size * (b + 1), len(sample))

            avg_batch_grad = reg_type.average_gradient(reg, sample, batch_order[b_begin:b_end]) + regularization.gradient(reg)

            reg = reg - learning_rate * avg_batch_grad

            avg_batch_loss = reg_type.average_loss(reg, sample, batch_order[b_begin:b_end]) + regularization.func(reg)

            avg_loss_pr = avg_loss
            avg_loss = decay * avg_batch_loss + (1 - decay) * avg_loss
        
        loss_trace.append(avg_loss)
        print('epoch {}, loss {}'.format(epoech, avg_loss))

        if (abs(avg_loss - avg_loss_pr) < eps):
            break
    
    finish_time = time() - start_time
    memory = tracemalloc.get_traced_memory()
    tracemalloc.stop()
    return reg, memory, finish_time, epoech, iterations, loss_trace

##########################################

class LinearRegressionModel(torch.nn.Module):
    def __init__(self):
        super(LinearRegressionModel, self).__init__()
        self.linear = torch.nn.Linear(1, 1)  # One in and one out
 
    def forward(self, x):
        y_pred = self.linear(x)
        return y_pred


def runOptimizer(model, optimizer, sample_X, sample_Y, batch_size=1, N=1000, eps=1e-2):
    tracemalloc.start()
    start_time = time()

    criterion = torch.nn.MSELoss()

    loss_pr = None
    loss = None

    for epoech in range(N):
        loss_pr = loss

        for b in range(0, max(1, math.floor(len(sample_X) / batch_size))):
            b_begin = batch_size * b
            b_end = min(batch_size * (b + 1), len(sample_X))

            data_x = Variable(torch.from_numpy(sample_X[b_begin:b_end]))
            data_y = Variable(torch.from_numpy(sample_Y[b_begin:b_end]))
            
            model_y = model(data_x)

            loss = criterion(model_y, data_y)
        
            optimizer.zero_grad()
            loss.backward()
            optimizer.step()

        #print('epoch {}, loss {}'.format(epoch, loss))

        if (loss and loss_pr and abs(loss - loss_pr) < eps):
            break

    finish_time = time() - start_time
    memory = tracemalloc.get_traced_memory()
    tracemalloc.stop()
    return memory, finish_time


#########################################

def poly(w):
    def f(x):
        return sum(x**i*w[i] for i in range(len(w)))
    return f

def func(x):
    return np.sin(2*x)

# Параметры
density = 10000
dots_count = 500
dist = 3
radius = 0.05

weights = np.array([4, 8, -8, -4])
f = func#poly(weights)

X = np.linspace(-dist, dist, density)
Y = np.array([f(x) for x in X])

sample_X, sample_Y = generate_sample(X, lambda x : f(x), dots_count, radius)
sample = [np.array([sample_X[i], sample_Y[i]]) for i in range(len(sample_X))]
reg_start = np.array([0, 0, 0, 0])

reg, memory_u, time_u, epoech, iter, loss = stochastic_gradient_descent(reg_start, sample, PolynomialRegression, LearningRate.constant(5e-3), L2(-1), batch_size=100, eps=1e-4)

Y_found = np.array([poly(reg)(x) for x in X])

#print(LinearRegression.average_loss(weights, sample, np.arange(0, len(sample))))
print(PolynomialRegression.average_loss(reg, sample, np.arange(0, len(sample))))

print_sample(X, Y, Y_found, sample_X, sample_Y)


lib_optimizers = [
    lambda params: torch.optim.SGD(params=params, lr=1e-2),
    lambda params: torch.optim.SGD(params=params, lr=1e-2, momentum=0.9),
    lambda params: torch.optim.SGD(params=params, lr=1e-2, momentum=0.9, nesterov=True),
    lambda params: torch.optim.Adagrad(params=params, lr=1000),
    lambda params: torch.optim.RMSprop(params=params, lr=10),
    lambda params: torch.optim.Adam(params=params, lr=10)
]

import docx
doc = docx.Document()

"""batch_step = math.floor(dots_count / 100)
batches = [1]
for batch in range(batch_step, dots_count + batch_step, batch_step): batches.append(batch)

table1 = doc.add_table(rows=len(batches), cols=5)
table1.style = 'Table Grid'

for i in range(len(batches)):
    reg, memory_u, time_u, epoech, iter, loss = stochastic_gradient_descent(reg_start, sample, LinearRegression, LearningRate.constant(1e-1), L2(0), batch_size=batches[i], eps=1e-4)
    
    table1.cell(i, 0).text = str(batches[i])
    table1.cell(i, 1).text = str(epoech)
    table1.cell(i, 2).text = str(iter)
    table1.cell(i, 3).text = str(round(time_u, 5))
    table1.cell(i, 4).text = str(memory_u[1])

    print(list(map(lambda x : round(x, 3), reg)), memory_u[1], time_u)

doc.save("test1.docx")"""

"""table2 = doc.add_table(rows=50, cols=4)
table2.style = 'Table Grid'

for epoeches in range(1, 51):
    reg, memory_u, time_u, epoech, iter, loss_trace_const = stochastic_gradient_descent(reg_start, sample, LinearRegression, LearningRate.constant(3e-1), L2(0), batch_size=dots_count, N=epoeches, eps=1e-4)
    reg, memory_u, time_u, epoech, iter, loss_trace_exp = stochastic_gradient_descent(reg_start, sample, LinearRegression, LearningRate.exponential(3e-1, 0.05), L2(0), batch_size=dots_count, N=epoeches, eps=1e-4)
    reg, memory_u, time_u, epoech, iter, loss_trace_step = stochastic_gradient_descent(reg_start, sample, LinearRegression, LearningRate.step_based(3e-1, 0.9, 5), L2(0), batch_size=dots_count, N=epoeches, eps=1e-4)

    table2.cell(epoech - 1, 0).text = str(epoech)
    table2.cell(epoech - 1, 1).text = str(round(loss_trace_const[-1], 3))
    table2.cell(epoech - 1, 2).text = str(round(loss_trace_exp[-1], 3))
    table2.cell(epoech - 1, 3).text = str(round(loss_trace_step[-1], 3))

    print(loss_trace_const[-1], loss_trace_exp[-1], loss_trace_step[-1])

doc.save("test2.docx")"""

"""table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Table Grid'

for optimizer_i in range(len(lib_optimizers)):
    model = LinearRegressionModel()
    memory_u, time_u = runOptimizer(model, lib_optimizers[optimizer_i](model.parameters()), sample_X.reshape(-1, 1), sample_Y.reshape(-1, 1), batch_size=10)
    
    # y = ax + b
    b = model(Variable(torch.Tensor([[0.0]]))).data.numpy()[0][0]
    a = model(Variable(torch.Tensor([[1.0]]))).data.numpy()[0][0] - b

    loss = LinearRegression.average_loss([b, a], sample, range(len(sample)))

    table3.cell(optimizer_i, 0).text = str(round(time_u, 5))
    table3.cell(optimizer_i, 1).text = str(memory_u[1])
    table3.cell(optimizer_i, 2).text = str(round(loss, 3))

    print(b, a, memory_u[1], time_u, loss)

doc.save("test3.docx")"""
